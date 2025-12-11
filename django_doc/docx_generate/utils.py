from docxtpl import DocxTemplate
from datetime import datetime
import os

def generate_document(template_path, context: dict):
    doc = DocxTemplate(template_path)
    doc.render(context)
    filename = f"generated_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join("media/generated", filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

from docx import Document
from zipfile import ZipFile
import io
import re
from lxml import etree

def extract_text_from_xml(docx_bytes):
    """Извлекает текст из всех <w:t> блоков .docx-файла (XML внутри архива)."""
    with ZipFile(io.BytesIO(docx_bytes)) as docx:
        xml_content = docx.read("word/document.xml")
        tree = etree.fromstring(xml_content)
        texts = tree.xpath("//w:t", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        return "\n".join(t.text for t in texts if t.text)


def extract_template_variables(docx_bytes):
    """Находит все {{ variable }} и {{ obj.field }} переменные."""
    xml = extract_text_from_xml(docx_bytes)
    return re.findall(r"{{\s*(.*?)\s*}}", xml)


def parse_template_structure(template_bytes):
    """Извлекает структуру шаблона: скалярные поля и таблицы."""
    doc = Document(io.BytesIO(template_bytes))
    full_text = " ".join([p.text for p in doc.paragraphs if p.text.strip()])
    variables = extract_template_variables(template_bytes)

    scalar_fields = [v for v in variables if '.' not in v]
    table_fields = {}

    table_pattern = r"{%\s*(?:tr\s+)?for\s+(\w+)\s+in\s+(\w+)\s*%}(.*?){%\s*endfor\s*%}"
    for match in re.finditer(table_pattern, full_text, re.DOTALL):
        loop_var = match.group(1)
        list_name = match.group(2)
        table_vars = [v for v in variables if v.startswith(f"{loop_var}.")]
        table_fields[list_name] = [v.replace(f"{loop_var}.", "") for v in table_vars]

    return {
        'scalar_fields': scalar_fields,
        'table_fields': table_fields,
        'full_text': full_text
    }



def reverse_docxtpl(template_bytes, filled_bytes):
    """Извлекает данные из документа, основываясь на шаблоне DocxTemplate."""
    filled_text = extract_text_from_xml(filled_bytes)
    doc = Document(io.BytesIO(filled_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    template_info = parse_template_structure(template_bytes)

    context = {"data": {}}

    # Извлекаем скалярные значения
    for field in template_info['scalar_fields']:
        context["data"][field] = ""
        for para in paragraphs:
            if field in ['report_date', 'contract_date'] and re.search(r"\d{4}-\d{2}-\d{2}", para):
                date_match = re.search(r"(\d{4}-\d{2}-\d{2})", para)
                if date_match:
                    context["data"][field] = date_match.group(1)
                    break
            else:
                words = para.split()
                for i, word in enumerate(words):
                    if word.strip() and not re.match(r"^\W+$", word):
                        context["data"][field] = word.strip()
                        break
                if context["data"][field]:
                    break

    # Извлекаем таблицы
    for list_name, fields in template_info['table_fields'].items():
        context["data"][list_name] = []
        table_content = filled_text.split('\n')
        for line in table_content:
            line = line.strip()
            if not line:
                continue
            values = line.split()
            if len(values) >= len(fields):
                item = {}
                for i, field in enumerate(fields):
                    if field in ['salary', 'price', 'amount'] and re.match(r"^\d+$", values[i]):
                        item[field] = values[i]
                    elif field in ['name', 'position', 'description']:
                        text = values[i]
                        j = i + 1
                        while j < len(values) and not re.match(r"^\d+$", values[j]):
                            text += " " + values[j]
                            j += 1
                        item[field] = text
                    else:
                        item[field] = values[i]
                context["data"][list_name].append(item)
            else:
                break

    return context
